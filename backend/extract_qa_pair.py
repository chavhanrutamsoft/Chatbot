# #!/usr/bin/env python3
# """
# FINAL PRODUCTION EXTRACTOR
# - Rule-based section split
# - LLM-assisted understanding (micro task)
# - Multi-view chunk generation
# """

# import json
# import os
# import re
# import time
# from pathlib import Path
# from typing import List, Dict
# from docx import Document
# from openai import OpenAI
# from dotenv import load_dotenv
# from concurrent.futures import ThreadPoolExecutor, as_completed
# from tqdm import tqdm

# # ---------------- CONFIG ----------------

# PROJECT_ROOT = Path(__file__).parent.parent
# DATA_DIR = PROJECT_ROOT / "data"

# # Load .env file from project root
# load_dotenv(PROJECT_ROOT / ".env")

# INPUT_DOCX = DATA_DIR / "1.docx"
# OUTPUT_JSON = DATA_DIR / "final_chunks.json"    

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# MODEL = "gpt-4o"
# MAX_WORKERS = int(os.getenv("MAX_WORKERS", "5"))  # Parallel API calls
# MAX_RETRIES = 3
# REQUEST_TIMEOUT = 30  # seconds

# if not OPENAI_API_KEY:
#     raise ValueError("OPENAI_API_KEY environment variable is not set. Please set it in your .env file or environment.")

# # Create client with timeout
# client = OpenAI(api_key=OPENAI_API_KEY, timeout=REQUEST_TIMEOUT)

# # ---------------- HELPERS ----------------

# def clean(text: str) -> str:
#     text = re.sub(r'\s+', ' ', text)
#     return text.strip()

# def is_heading(text: str) -> bool:
#     return (
#         text.isupper()
#         or text.endswith(":")
#         or re.match(r"^\d+(\.\d+)*\s", text)
#         or len(text) < 80 and text.istitle()
#     )

# # ---------------- LLM MICRO TASK ----------------

# def analyze_section_with_llm(heading: str, content: str, retry_count: int = 0) -> Dict:
#     """Analyze section with LLM, with retry logic and timeout."""
#     prompt = f"""
# You are analyzing a software manual section.

# HEADING:
# {heading}

# CONTENT:
# {content}

# TASK:
# Return JSON ONLY:
# {{
#   "term": "main concept name",
#   "definition": "1-2 line definition or empty",
#   "usage": "why/where it is used or empty",
#   "has_procedure": true/false
# }}
# """
#     try:
#         res = client.chat.completions.create(
#             model=MODEL,
#             temperature=0,
#             messages=[{"role": "user", "content": prompt}],
#         )

#         try:
#             return json.loads(res.choices[0].message.content)
#         except Exception:
#             return {
#                 "term": heading.strip(":"),
#                 "definition": "",
#                 "usage": "",
#                 "has_procedure": False
#             }
#     except Exception as e:
#         if retry_count < MAX_RETRIES:
#             # Exponential backoff
#             wait_time = 2 ** retry_count
#             time.sleep(wait_time)
#             return analyze_section_with_llm(heading, content, retry_count + 1)
#         else:
#             # Return fallback after max retries
#             print(f"Warning: Failed to analyze section '{heading[:50]}...' after {MAX_RETRIES} retries. Using fallback.")
#             return {
#                 "term": heading.strip(":"),
#                 "definition": "",
#                 "usage": "",
#                 "has_procedure": False
#             }

# # ---------------- STEP EXTRACTION ----------------

# def extract_steps(text: str) -> List[str]:
#     steps = []
#     for m in re.finditer(r"\d+[\).\s]+(.+?)(?=\d+[\).\s]|$)", text):
#         step = clean(m.group(1))
#         if len(step) > 10:
#             steps.append(step)
#     return steps

# # ---------------- MAIN EXTRACTION ----------------

# def extract_chunks(docx_path: Path) -> List[Dict]:
#     doc = Document(docx_path)
    
#     # First pass: collect all sections
#     sections = []
#     current_heading = None
#     buffer = []

#     for p in doc.paragraphs:
#         text = clean(p.text)
#         if not text:
#             continue

#         if is_heading(text):
#             if current_heading and buffer:
#                 sections.append((current_heading, " ".join(buffer)))
#             current_heading = text
#             buffer = []
#         else:
#             buffer.append(text)

#     if current_heading and buffer:
#         sections.append((current_heading, " ".join(buffer)))

#     print(f"Found {len(sections)} sections. Processing with {MAX_WORKERS} parallel workers...")
    
#     # Second pass: process sections in parallel
#     chunks = []
    
#     with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
#         # Submit all tasks
#         future_to_section = {
#             executor.submit(process_section_parallel, heading, content): (heading, content)
#             for heading, content in sections
#         }
        
#         # Process completed tasks with progress bar
#         for future in tqdm(as_completed(future_to_section), total=len(sections), desc="Processing sections"):
#             heading, content = future_to_section[future]
#             try:
#                 section_chunks = future.result()
#                 chunks.extend(section_chunks)
#             except Exception as e:
#                 print(f"Error processing section '{heading[:50]}...': {e}")
#                 # Add fallback chunk
#                 chunks.append({
#                     "type": "raw",
#                     "text": f"{heading} information:\n{content}"
#                 })

#     return chunks




# def process_section_parallel(heading: str, content: str) -> List[Dict]:
#     """Process a single section and return chunks (for parallel processing)."""
#     llm = analyze_section_with_llm(heading, content)
#     term = llm.get("term") or heading.strip(":")
#     definition = llm.get("definition", "")
#     usage = llm.get("usage", "")
#     has_proc = llm.get("has_procedure", False)

#     section_chunks = []

#     # 1️⃣ Definition
#     if definition:
#         section_chunks.append({
#             "type": "definition",
#             "text": f"Definition of {term}. {definition}"
#         })
#         section_chunks.append({
#             "type": "meaning",
#             "text": f"{term} means {definition}"
#         })
#         section_chunks.append({
#             "type": "qa",
#             "text": f"Q: What is {term}?\nA: {definition}"
#         })

#     # 2️⃣ Usage / Explanation
#     if usage:
#         section_chunks.append({
#             "type": "usage",
#             "text": f"{term} is used for {usage}"
#         })
#         section_chunks.append({
#             "type": "qa",
#             "text": f"Q: Why is {term} used?\nA: {usage}"
#         })

#     # 3️⃣ Procedure
#     if has_proc:
#         steps = extract_steps(content)
#         if steps:
#             step_text = "\n".join(f"{i+1}. {s}" for i, s in enumerate(steps))
#             section_chunks.append({
#                 "type": "procedure",
#                 "text": f"How to work with {term}:\n{step_text}"
#             })

#     # 4️⃣ Raw fallback (safety net)
#     section_chunks.append({
#         "type": "raw",
#         "text": f"{term} information:\n{content}"
#     })
    
#     return section_chunks

# # ---------------- RUN ----------------

# if __name__ == "__main__":
#     import sys
#     # Set UTF-8 encoding for Windows console
#     if sys.platform == "win32":
#         sys.stdout.reconfigure(encoding='utf-8')
    
#     print("Running FINAL PRODUCTION EXTRACTOR")
#     chunks = extract_chunks(INPUT_DOCX)

#     with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
#         json.dump(chunks, f, indent=2, ensure_ascii=False)

#     print(f"Extraction complete")
#     print(f"Total chunks created: {len(chunks)}")
#     print(f"Saved to: {OUTPUT_JSON}")



#!/usr/bin/env python3
"""
FINAL PRODUCTION EXTRACTOR (RAG-OPTIMIZED)
- Rule-based section detection
- Lossless DOCX → JSON conversion
- Step-level procedure extraction
- NO LLM during ingestion (CRITICAL)
"""

import json
import re
from pathlib import Path
from typing import List, Dict
from docx import Document

# ---------------- CONFIG ----------------

PROJECT_ROOT = Path(__file__).parent.parent
DATA_DIR = PROJECT_ROOT / "data"

INPUT_DOCX = DATA_DIR / "QuotePlan.docx"
OUTPUT_JSON = DATA_DIR / "final_chunks.json"

# ---------------- HELPERS ----------------

def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()

def is_heading(text: str) -> bool:
    return (
        text.isupper()
        or text.endswith(":")
        or re.match(r"^\d+(\.\d+)*\s+", text)
        or (len(text) < 80 and text.istitle())
    )

STEP_REGEX = re.compile(
    r"(?:^|\n)(\d+[\.\)]\s+.+?)(?=\n\d+[\.\)]|\Z)",
    re.DOTALL
)

# ---------------- EXTRACTION ----------------

def extract_chunks(docx_path: Path) -> List[Dict]:
    doc = Document(docx_path)
    chunks: List[Dict] = []

    current_section = "General"
    buffer: List[str] = []
    chunk_id = 1

    def flush_paragraphs():
        nonlocal chunk_id
        if not buffer:
            return

        text_block = "\n".join(buffer)

        # Extract steps if present
        steps = STEP_REGEX.findall(text_block)
        if steps:
            for order, step in enumerate(steps, start=1):
                chunks.append({
                    "id": f"chunk_{chunk_id}",
                    "section": current_section,
                    "type": "procedure_step",
                    "order": order,
                    "text": clean(step)
                })
                chunk_id += 1
        else:
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "section": current_section,
                "type": "paragraph",
                "text": clean(text_block)
            })
            chunk_id += 1

        buffer.clear()

    for p in doc.paragraphs:
        text = clean(p.text)
        if not text:
            continue

        if is_heading(text):
            flush_paragraphs()
            current_section = text.rstrip(":")
        else:
            buffer.append(text)

    flush_paragraphs()
    return chunks

# ---------------- RUN ----------------

if __name__ == "__main__":
    print("Running RAG-optimized extractor...")
    chunks = extract_chunks(INPUT_DOCX)

    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(chunks, f, indent=2, ensure_ascii=False)

    print("Extraction complete")
    print(f"Total chunks created: {len(chunks)}")
    print(f"Saved to: {OUTPUT_JSON}")
